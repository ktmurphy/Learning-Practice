import docx

doc = docx.Document("demo.docx")
doc.paragraphs[0].style = "Normal"
doc.paragraphs[1].runs[0].style(
    "Quote", "char"
)  # this code does not work, need to look up
doc.paragraphs[1].runs[1].underline = True
doc.save("restyled.docx")
