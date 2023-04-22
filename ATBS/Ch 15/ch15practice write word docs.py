import docx

doc = docx.Document()
doc.add_paragraph("Hello, World")
paraObj1 = doc.add_paragraph("This is a second paragraph")
paraObj2 = doc.add_paragraph("This is yet another paragraph")
paraObj1.add_run("This sentence is being added to the end of the second paragraph")
doc.add_paragraph("Hello World", "Title")
doc.add_heading("Header 0", 0)
doc.add_heading("Header 1", 1)
doc.add_heading("Header 2", 2)
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph("This is on the second page")
doc.add_picture("zophie.png", width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.save("helloworldpractice.docx")
